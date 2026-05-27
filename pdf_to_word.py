"""
PDF to Word Konvertor v1.0
Haqiqiy ishlaydigan PDF -> Word (.docx) konvertori
"""

import sys
import os
import threading
import tkinter as tk
from tkinter import ttk, filedialog, messagebox


def convert_pdf_to_word(pdf_path, docx_path=None, progress_cb=None):
    from pdfminer.high_level import extract_pages
    from pdfminer.layout import LTTextBox, LTTextLine, LTChar, LTAnno
    from docx import Document
    from docx.shared import Pt, Inches

    if not docx_path:
        base = os.path.splitext(pdf_path)[0]
        docx_path = base + ".docx"

    doc = Document()
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    pages = list(extract_pages(pdf_path))
    total_pages = len(pages)

    for page_num, page_layout in enumerate(pages):
        if progress_cb:
            progress_cb(int((page_num / total_pages) * 90) + 5)

        page_texts = []

        for element in page_layout:
            if isinstance(element, LTTextBox):
                box_text = ""
                font_sizes = []
                is_bold = False

                for line in element:
                    if isinstance(line, LTTextLine):
                        line_text = ""
                        for char in line:
                            if isinstance(char, LTChar):
                                line_text += char.get_text()
                                font_sizes.append(char.size)
                                fname = char.fontname.lower() if char.fontname else ""
                                if "bold" in fname or "heavy" in fname:
                                    is_bold = True
                            elif isinstance(char, LTAnno):
                                line_text += char.get_text()
                        box_text += line_text

                if box_text.strip():
                    avg_size = sum(font_sizes) / len(font_sizes) if font_sizes else 12
                    page_texts.append({
                        "text": box_text,
                        "size": avg_size,
                        "bold": is_bold,
                        "y": element.y1
                    })

        page_texts.sort(key=lambda x: -x["y"])

        for item in page_texts:
            text = item["text"].strip()
            if not text:
                continue
            para = doc.add_paragraph()
            if item["size"] >= 16:
                para.style = "Heading 1"
                para.add_run(text)
            elif item["size"] >= 13:
                para.style = "Heading 2"
                para.add_run(text)
            else:
                run = para.add_run(text)
                run.font.size = Pt(max(8, min(14, item["size"])))
                run.bold = item["bold"]

        if page_num < total_pages - 1:
            doc.add_page_break()

    doc.save(docx_path)
    if progress_cb:
        progress_cb(100)
    return docx_path


def run_gui():
    root = tk.Tk()
    root.title("PDF → Word Konvertor")
    root.geometry("540x400")
    root.resizable(False, False)
    root.configure(bg="#f7fafc")

    # Header
    header = tk.Frame(root, bg="#4c51bf", height=70)
    header.pack(fill="x")
    header.pack_propagate(False)
    tk.Label(header, text="PDF  →  Word Konvertor",
             font=("Segoe UI", 16, "bold"), bg="#4c51bf", fg="white").pack(expand=True)

    # Body
    body = tk.Frame(root, bg="#f7fafc", padx=30, pady=20)
    body.pack(fill="both", expand=True)

    # PDF row
    tk.Label(body, text="PDF fayl:", font=("Segoe UI", 10, "bold"),
             bg="#f7fafc", fg="#2d3748").grid(row=0, column=0, sticky="w", pady=8)

    pdf_var = tk.StringVar()
    tk.Entry(body, textvariable=pdf_var, font=("Segoe UI", 9),
             width=36, state="readonly", bg="white",
             relief="solid", bd=1).grid(row=0, column=1, padx=(5, 5), pady=8)

    def browse_pdf():
        f = filedialog.askopenfilename(
            title="PDF fayl tanlang",
            filetypes=[("PDF fayllari", "*.pdf"), ("Barcha fayllar", "*.*")])
        if f:
            pdf_var.set(f)
            out_var.set(os.path.splitext(f)[0] + ".docx")

    tk.Button(body, text="Tanlash", command=browse_pdf,
              bg="#ebf8ff", fg="#2b6cb0", font=("Segoe UI", 9),
              relief="solid", bd=1, padx=8, cursor="hand2").grid(row=0, column=2, pady=8)

    # Output row
    tk.Label(body, text="Saqlash joyi:", font=("Segoe UI", 10, "bold"),
             bg="#f7fafc", fg="#2d3748").grid(row=1, column=0, sticky="w", pady=8)

    out_var = tk.StringVar()
    tk.Entry(body, textvariable=out_var, font=("Segoe UI", 9),
             width=36, bg="white", relief="solid", bd=1).grid(row=1, column=1, padx=(5, 5), pady=8)

    def browse_out():
        f = filedialog.asksaveasfilename(
            title="Word faylni saqlash",
            defaultextension=".docx",
            filetypes=[("Word fayllari", "*.docx")])
        if f:
            out_var.set(f)

    tk.Button(body, text="Tanlash", command=browse_out,
              bg="#f0fff4", fg="#276749", font=("Segoe UI", 9),
              relief="solid", bd=1, padx=8, cursor="hand2").grid(row=1, column=2, pady=8)

    # Progress
    tk.Label(body, text="", bg="#f7fafc").grid(row=2, column=0, columnspan=3, pady=4)

    progress_var = tk.IntVar()
    style = ttk.Style()
    style.theme_use("default")
    style.configure("blue.Horizontal.TProgressbar",
                    background="#4c51bf", troughcolor="#e2e8f0")
    progress = ttk.Progressbar(body, variable=progress_var, maximum=100,
                                length=440, style="blue.Horizontal.TProgressbar")
    progress.grid(row=3, column=0, columnspan=3, pady=5)

    status_var = tk.StringVar(value="PDF fayl tanlang va boshlang")
    status_lbl = tk.Label(body, textvariable=status_var,
                          font=("Segoe UI", 9), bg="#f7fafc", fg="#718096")
    status_lbl.grid(row=4, column=0, columnspan=3, pady=5)

    def do_convert_thread():
        pdf = pdf_var.get()
        out = out_var.get()

        if not pdf:
            messagebox.showwarning("Ogohlantirish", "Iltimos, PDF fayl tanlang!")
            btn.config(state="normal", text="  Konvertatsiya qilish  ")
            return
        if not out:
            messagebox.showwarning("Ogohlantirish", "Saqlash joyi tanlanmagan!")
            btn.config(state="normal", text="  Konvertatsiya qilish  ")
            return
        if not os.path.exists(pdf):
            messagebox.showerror("Xato", f"Fayl topilmadi:\n{pdf}")
            btn.config(state="normal", text="  Konvertatsiya qilish  ")
            return

        def upd(val):
            progress_var.set(val)
            status_var.set(f"Ishlayapti... {val}%")

        try:
            result = convert_pdf_to_word(pdf, out, progress_cb=upd)
            status_var.set("Muvaffaqiyatli konvertatsiya qilindi!")
            status_lbl.config(fg="#38a169")
            if messagebox.askyesno("Tayyor!", f"Word fayl saqlandi:\n{result}\n\nFaylni ochmoqchimisiz?"):
                os.startfile(result)
        except Exception as e:
            status_var.set("Xatolik yuz berdi!")
            status_lbl.config(fg="#e53e3e")
            messagebox.showerror("Xato", f"Xatolik:\n{str(e)}")
        finally:
            btn.config(state="normal", text="  Konvertatsiya qilish  ")

    def do_convert():
        btn.config(state="disabled", text="  Ishlayapti...  ")
        status_lbl.config(fg="#4a5568")
        progress_var.set(0)
        threading.Thread(target=do_convert_thread, daemon=True).start()

    btn = tk.Button(body, text="  Konvertatsiya qilish  ",
                    command=do_convert,
                    bg="#4c51bf", fg="white",
                    font=("Segoe UI", 11, "bold"),
                    relief="flat", padx=20, pady=10,
                    cursor="hand2")
    btn.grid(row=5, column=0, columnspan=3, pady=20)

    root.mainloop()


if __name__ == "__main__":
    if len(sys.argv) >= 2:
        print(f"Konvertatsiya: {sys.argv[1]}")
        def show(v): print(f"\r{'#'*(v//5)}{' '*(20-v//5)} {v}%", end="")
        r = convert_pdf_to_word(sys.argv[1], sys.argv[2] if len(sys.argv) >= 3 else None, show)
        print(f"\nTayyor: {r}")
    else:
        run_gui()
